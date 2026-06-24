"""
多格式文档加载器
支持：PDF / Word / Markdown / 网页
统一接口，自动识别文件类型
"""

import re
from pathlib import Path
from typing import List, Optional
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup
from langchain_core.documents import Document


class DocumentLoader:
    """统一的文档加载器，根据文件类型自动分派"""

    LOADER_MAP = {
        ".pdf": "_load_pdf",
        ".docx": "_load_docx",
        ".doc": "_load_docx",
        ".md": "_load_markdown",
        ".txt": "_load_text",
    }

    def __init__(self, encoding: str = "utf-8"):
        self.encoding = encoding

    def load(self, source: str) -> List[Document]:
        if self._is_url(source):
            return self._load_web(source)
        return self.load_file(source)

    def load_file(self, file_path: str) -> List[Document]:
        fp = Path(file_path)
        if not fp.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        suffix = fp.suffix.lower()
        if suffix not in self.LOADER_MAP:
            raise ValueError(f"Unsupported format: {suffix}. Supported: {list(self.LOADER_MAP.keys())}")
        loader_method = getattr(self, self.LOADER_MAP[suffix])
        return loader_method(file_path)

    @staticmethod
    def _is_url(source: str) -> bool:
        parsed = urlparse(source)
        return parsed.scheme in ("http", "https")

    def _load_pdf(self, file_path: str) -> List[Document]:
        # PyPDF2 was renamed to pypdf (v4+). Support both for compatibility.
        try:
            from pypdf import PdfReader
        except ImportError:
            from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        docs = []
        filename = Path(file_path).name
        skipped = 0
        for i, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if not text or not text.strip():
                continue
            text = text.strip()
            # Skip table-of-contents pages (high dot density + short lines)
            if self._is_toc_page(text, i, len(reader.pages)):
                skipped += 1
                continue
            # Skip acknowledgment pages
            if self._is_acknowledgment(text):
                skipped += 1
                continue
            docs.append(Document(
                page_content=text,
                metadata={
                    "source": file_path,
                    "filename": filename,
                    "file_type": "pdf",
                    "page": i,
                    "total_pages": len(reader.pages),
                }
            ))
        if skipped:
            print(f"[LOADER] Skipped {skipped} low-value pages (TOC/acknowledgment)")
        return docs

    @staticmethod
    def _is_toc_page(text: str, page_num: int, total_pages: int) -> bool:
        """Detect if a page is a table of contents."""
        # TOC pages are usually early in the document
        if page_num > total_pages * 0.15:
            return False
        lines = text.split(chr(10))
        if len(lines) < 3:
            return False
        # TOC lines have lots of dots/periods connecting title to page number
        dot_lines = sum(1 for l in lines if l.count('.') >= 5 or l.count('..') >= 2)
        # Check for TOC keywords
        first_100 = text[:100]
        toc_keywords = [chr(0x76ee)+chr(0x5f55), '目录', 'Contents', 'CONTENTS']
        has_toc_kw = any(kw in first_100 for kw in toc_keywords)
        return has_toc_kw or (dot_lines >= len(lines) * 0.4 and len(lines) >= 5)

    @staticmethod
    def _is_acknowledgment(text: str) -> bool:
        """Detect if a page is an acknowledgment section."""
        first_line = text.split(chr(10))[0].strip() if text else ""
        ack_keywords = [chr(0x81f4)+chr(0x8c22), '致谢', 'Acknowledgement', 'ACKNOWLEDGEMENT']
        return any(kw in first_line for kw in ack_keywords)

    def _load_docx(self, file_path: str) -> List[Document]:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        filename = Path(file_path).name
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        full_text = "\n".join(paragraphs)
        if not full_text.strip():
            return []
        return [Document(
            page_content=full_text,
            metadata={
                "source": file_path,
                "filename": filename,
                "file_type": "docx",
            }
        )]

    def _load_markdown(self, file_path: str) -> List[Document]:
        raw_text = Path(file_path).read_text(encoding=self.encoding)
        filename = Path(file_path).name
        if not raw_text.strip():
            return []
        return [Document(
            page_content=raw_text,
            metadata={
                "source": file_path,
                "filename": filename,
                "file_type": "markdown",
            }
        )]

    def _load_text(self, file_path: str) -> List[Document]:
        return self._load_markdown(file_path)

    def _load_web(self, url: str) -> List[Document]:
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
        resp.encoding = resp.apparent_encoding or self.encoding
        soup = BeautifulSoup(resp.text, "lxml")
        for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()
        title = ""
        if soup.title:
            title = soup.title.get_text(strip=True)
        text = soup.get_text(separator="\n")
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        clean_text = "\n".join(lines)
        if not clean_text:
            return []
        return [Document(
            page_content=clean_text,
            metadata={
                "source": url,
                "filename": title or url,
                "file_type": "web",
                "title": title,
            }
        )]


def load_document(source: str, encoding: str = "utf-8") -> List[Document]:
    """One-liner to load a document."""
    return DocumentLoader(encoding=encoding).load(source)


def load_documents(sources: List[str], encoding: str = "utf-8") -> List[Document]:
    """Batch load multiple documents."""
    loader = DocumentLoader(encoding=encoding)
    all_docs = []
    for src in sources:
        try:
            all_docs.extend(loader.load(src))
        except Exception as e:
            print(f"[WARN] Failed to load {src}: {e}")
    return all_docs
